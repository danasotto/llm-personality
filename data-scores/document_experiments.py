from io import BytesIO
import os
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import pandas as pd
import re
import seaborn as sns
import numpy as np
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def get_model_stats(df, modelname):
    

        
    # Group by post_id and sum word counts per post
    df = df.drop_duplicates(subset=["post_id", "word_count", "generated_comment"])
    word_count_per_post = df.groupby("post_id")["word_count"].sum()


    # Get the max, min, and mean of the total word count per post
    max_wc = word_count_per_post.max()
    min_wc = word_count_per_post.min()
    mean_wc = word_count_per_post.mean()

       
    stats_text = (
        f"Number of comments for {modelname}: {len(df)}\n"
        f"Mean comments word count: {df['word_count'].mean()}\n"
        f"Median comments word count: {df['word_count'].median()}\n"
        f"Max comments word count: {df['word_count'].max()}\n"
        f"Min comments word count: {df['word_count'].min()}\n"
        f"Std comments word count: {df['word_count'].std()}\n"
        f"Max word count per post (model output length): {max_wc}\n"
        f"Min word count per post: {min_wc}\n"
        f"Mean word count per post: {mean_wc}\n"
    )
    return stats_text

def generate_trait_graphs(df, model_name):
    traits = ['ext', 'opn', 'neu', 'agr', 'con']
    graph_paths = []

    for trait in traits:
        plt.hist(df[trait], bins=100)
        plt.xlabel(trait)
        plt.ylabel("Count")
        plt.title(f"Distribution of {trait.upper()} - {model_name}")

        # Save the figure
        img_stream = BytesIO()
        plt.savefig(img_stream, format='png')
        plt.close()
        img_stream.seek(0)

        graph_paths.append(img_stream)

    return graph_paths

def create_subreddits_graph(df):
        # List of personality traits to analyze
    traits = ['ext', 'opn', 'neu', 'agr', 'con']
    trait_names = {
        'ext': 'Extraversion',
        'opn': 'Openness',
        'neu': 'Neuroticism',
        'agr': 'Agreeableness',
        'con': 'Conscientiousness'
    }

    # Create a figure
    plt.figure(figsize=(18, 12))

    # Create a 3x2 grid of subplots (one extra for the legend)
    gs = plt.GridSpec(3, 2, height_ratios=[1, 1, 0.2])

    # For each trait, plot top 5 and bottom 5 subreddits
    for i, trait in enumerate(traits):
        # Calculate position in the grid
        row = i // 2
        col = i % 2

        # Create subplot
        ax = plt.subplot(gs[row, col])

        # Group by subreddit and calculate mean trait value
        trait_by_subreddit = df.groupby('post_subreddit')[trait].mean().sort_values(ascending=False)

        # Get top 5 and bottom 5
        top5 = trait_by_subreddit.head(5)
        bottom5 = trait_by_subreddit.tail(5)

        # Combine them with a separator
        combined = pd.concat([top5, pd.Series([np.nan], index=['_______']), bottom5])

        # Create a color map: green for top 5, red for bottom 5
        colors = ['#2ecc71']*5 + ['white'] + ['#e74c3c']*5

        # Plot
        # bars = sns.barplot(x=combined.index, y=combined.values, palette=colors, ax=ax)
        bars = sns.barplot(x=combined.index, y=combined.values, ax=ax)
        for i, bar in enumerate(bars.patches):
            bar.set_color(colors[i])

        # Add value labels on top of each bar
        for j, bar in enumerate(bars.patches):
            if j != 5:  # Skip the separator
                ax.text(
                    bar.get_x() + bar.get_width()/2.,
                    bar.get_height() + 0.01,
                    f'{combined.values[j]:.3f}' if not np.isnan(combined.values[j]) else '',
                    ha="center", fontsize=9
                )

        # Customize the subplot
        ax.set_title(f'{trait_names[trait]}', fontsize=16)
        ax.set_xlabel('Subreddit', fontsize=12)
        ax.set_ylabel(f'Average Score', fontsize=12)
        ax.tick_params(axis='x', rotation=45)
        ax.grid(axis='y', linestyle='--', alpha=0.7)

        # Add a horizontal line for the overall mean
        overall_mean = df[trait].mean()
        ax.axhline(y=overall_mean, color='blue', linestyle='--', alpha=0.7)
        ax.text(
            0, overall_mean + 0.01, 
            f'Mean: {overall_mean:.3f}', 
            color='blue', fontsize=9
        )

    # Create a custom legend in the bottom subplot
    legend_ax = plt.subplot(gs[2, :])
    legend_ax.axis('off')
    legend_handles = [
        plt.Rectangle((0,0),1,1, color='#2ecc71'),
        plt.Rectangle((0,0),1,1, color='#e74c3c'),
        plt.Line2D([0], [0], color='blue', linestyle='--')
    ]

    plt.suptitle(f"Top and Bottom 5 Subreddits by Big Five ({model_name})", fontsize=20, y=0.98)
    plt.tight_layout(rect=[0, 0, 1, 0.96])

    img_stream = BytesIO()
    plt.savefig(img_stream, format='png')
    plt.close()
    img_stream.seek(0)

    return img_stream
    
# Function to add an experiment entry
def add_experiment(doc, number, model, prompt, stats_text, subreddit_graph, trait_graphs):
    # Heading
    doc.add_heading(f'Experiment #{number}: Model = {model}', level=1)
    
    # Prompt
    if ("Reddit" not in model_name):
        doc.add_paragraph("Prompt =")
        doc.add_paragraph(prompt).paragraph_format.left_indent = Pt(20)
    
    # Stats
    doc.add_paragraph("Stats =")
    doc.add_paragraph(stats_text).paragraph_format.left_indent = Pt(20)
    
    doc.add_paragraph("")  # Spacer line

    # Add the graph to the Word file
    for graph in trait_graphs:
        doc.add_picture(graph, width=Inches(5))
    
    # Add the graph to the Word file
    doc.add_picture(subreddit_graph, width=Inches(5))


# Initialize document
doc = Document()

# Add a title
doc.add_heading('Experiment Report', 0)

prompt = (
    "You are several Reddit users. "
    f"Generate {{num_long_comments}} comments, at least 100 words long, "
    "in response to the following Reddit post. Number each comment:\n\n"
    f"Post's title: {{title}}\n"
    f"Post content: {{text}}\n\n"
    "Only return the numbered list of comments."
)


# Path to your folder containing model CSVs
data_dir = 'C:\\Users\\ARIK_LAP\\Desktop\\Dana\\פרויקט גמר תואר שני\\reddit-data-with-code\\data-scores\\llm models with predictions'
output_filename = 'model_summary_stats.csv'
output_file = os.path.join(data_dir, output_filename)

# Create a dictionary to hold model name -> DataFrame
models = {}

subreddit_df = pd.read_csv('C:\\Users\\ARIK_LAP\\Desktop\\Dana\\פרויקט גמר תואר שני\\reddit-data-with-code\\reddit_posts_clean.csv')
subreddit_df = subreddit_df.drop_duplicates(subset='post_id')

# Loop through all CSV files in the directory
for filename in os.listdir(data_dir):

    if  re.search('comments', filename) and filename.endswith('.csv'):
        model_name =  filename.split("comments_")[1].split("_with")[0]
        file_path = os.path.join(data_dir, filename)
        df = pd.read_csv(file_path)

        if "post_subreddit" not in df.columns:
            df = df.merge(
                subreddit_df[['post_id', 'post_subreddit']], 
                on='post_id',
                how='left'  # Use 'left' to keep all rows from your original dataframe
            )

        models[model_name] = df

traits = ['ext', 'opn', 'neu', 'agr', 'con']
counter = 0

for model_name, df in models.items():
    counter += 1

    print (model_name + ":\n")
    # get model stats
    stats_text = get_model_stats(df, model_name)

    #create plots
    trait_graphs = generate_trait_graphs(df,model_name)

    #plot of subreddits
    subreddit_graph = create_subreddits_graph(df)

    # Add to document
    add_experiment(doc, counter, model_name, prompt, stats_text, subreddit_graph, trait_graphs)


# Save the report
doc.save(f"{data_dir}\\experiment_report.docx")